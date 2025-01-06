import streamlit as st
from docx import Document
from docx.shared import Pt
import os
import google.generativeai as genai
import io
from pypdf import PdfReader
import re
import yaml
from yaml.loader import SafeLoader
import time
from dotenv import load_dotenv
import json
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches

# --- Load Environment Variables ---
load_dotenv()

# --- Setup ---
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    st.error("Please set your Google API key in the .env file.")
    st.stop()

genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

# --- Caching ---
generation_cache = {}  # Simple dictionary for caching

# --- Helper Functions ---
# --- Helper Functions ---

def generate_roadmap(subject, syllabus_text, difficulty_level, temperature=0.7):
    """Generates a detailed roadmap from the syllabus."""
    prompt = f"""
You are an expert educator tasked with creating a detailed roadmap for the subject: "{subject}".

**Syllabus:** {syllabus_text}
Target Audience: {difficulty_level} level students

Your task is to generate a comprehensive roadmap that outlines the entire syllabus, divided into main topics, subtopics, and further sub-divisions if necessary. The output MUST STRICTLY ADHERE to the following hierarchical format and output nothing else:**

   T<number>: Main Topic Description (e.g., `T1: Introduction to Programming`)
       T<number>.<number>: Subtopic Description (e.g., `T1.1: Basic Data Types`)
           T<number>.<number>.<number>: Sub-subtopic Description (e.g., `T1.1.1: Integers and Floats`)
               T<number>.<number>.<number>.<number>: Further sub-division Description (if needed)

**Rules:**

1. **Hierarchical Format:** Use the exact hierarchical format specified above with "T" followed by numbers and dots.
2. **Topic and Subtopic Descriptions:** Each topic and subtopic MUST be followed by a colon (`:`) and a brief description on the SAME LINE.
3. **NO Extra Text:** Do not include any introductory text, explanations, or additional formatting beyond what is shown in the example structure.
4. **NO Aestericks:** Do not use any aestericks in the output.
5. **STRICT ADHERENCE:** The output must strictly follow these rules. Any deviations from this format will make the roadmap unusable.

**Consider the principles of chunking and scaffolding when organizing the outline. Suggest a logical sequence for the topics (Linear, Spiral, or Modular) in a separate line at the beginning, using the format:** `Sequence: <Sequence Type>` (e.g., `Sequence: Linear`).
    """
    try:
        if prompt in generation_cache:
            st.success("Roadmap found in cache!")
            return generation_cache[prompt]
        with st.spinner("Generating roadmap..."):
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=temperature,
                )
            )
        st.success("Roadmap generated successfully!")
        generation_cache[prompt] = response.text.strip()
        return response.text.strip()
    except Exception as e:
        st.error(f"Error generating roadmap: {e}")
        print(f"Error generating roadmap: {e}")  # Debugging: Print error to console
        return ""

def parse_roadmap(roadmap_text):
    """
    Parses a roadmap string into a structured dictionary using regular expressions.
    """
    roadmap = {"topics": []}
    lines = roadmap_text.split("\n")

    main_topic_re = r"^T(\d+):\s*(.+)$"
    subtopic_re = r"^T(\d+)\.(\d+):\s*(.+)$"
    subsubtopic_re = r"^T(\d+)\.(\d+)\.(\d+):\s*(.+)$"
    subsubsubtopic_re = r"^T(\d+)\.(\d+)\.(\d+)\.(\d+):\s*(.+)$"

    current_topic = None
    current_subtopic = None
    current_subsubtopic = None
    current_subsubsubtopic = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        main_match = re.match(main_topic_re, line)
        sub_match = re.match(subtopic_re, line)
        subsub_match = re.match(subsubtopic_re, line)
        subsubsub_match = re.match(subsubsubtopic_re, line)

        if main_match:
            topic_num = int(main_match.group(1))
            topic_desc = main_match.group(2)
            current_topic = {
                "id": f"T{topic_num}",
                "description": topic_desc,
                "subtopics": [],
            }
            roadmap["topics"].append(current_topic)
            current_subtopic = None
            current_subsubtopic = None
            current_subsubsubtopic = None
        elif sub_match:
            topic_num = int(sub_match.group(1))
            subtopic_num = int(sub_match.group(2))
            subtopic_desc = sub_match.group(3)
            current_subtopic = {
                "id": f"T{topic_num}.{subtopic_num}",
                "description": subtopic_desc,
                "subsubtopics": [],
            }
            if current_topic:
                current_topic["subtopics"].append(current_subtopic)
            current_subsubtopic = None
            current_subsubsubtopic = None
        elif subsub_match:
            topic_num = int(subsub_match.group(1))
            subtopic_num = int(subsub_match.group(2))
            subsubtopic_num = int(subsub_match.group(3))
            subsubtopic_desc = subsub_match.group(4)
            current_subsubtopic = {
                "id": f"T{topic_num}.{subtopic_num}.{subsubtopic_num}",
                "description": subsubtopic_desc,
                "subsubsubtopics": [],
            }
            if current_subtopic:
                current_subtopic["subsubtopics"].append(current_subsubtopic)
            current_subsubsubtopic = None
        elif subsubsub_match:
            topic_num = int(subsubsub_match.group(1))
            subtopic_num = int(subsubsub_match.group(2))
            subsubtopic_num = int(subsubsub_match.group(3))
            subsubsubtopic_num = int(subsubsub_match.group(4))
            subsubsubtopic_desc = subsubsub_match.group(5)
            current_subsubsubtopic = {
                "id": f"T{topic_num}.{subtopic_num}.{subsubtopic_num}.{subsubsubtopic_num}",
                "description": subsubsubtopic_desc,
                "details": [],
            }
            if current_subsubtopic:
                current_subsubtopic["subsubsubtopics"].append(current_subsubsubtopic)
        else:
            print(f"Warning: Could not parse line: {line}")

    return roadmap

def build_prompt_with_hierarchy(subject, difficulty_level, topic_data, parent_topics_content=None, depth=1):
    """
    Builds a highly optimized prompt for generating lesson plan content, including hierarchical context, specific instructions to prevent repetition, and incorporating advanced learning strategies.
    """
    topic_details = f"**Topic:** {topic_data['id']}: {topic_data['description']}\n"

    prompt = f"""
You are an expert educator creating a detailed lesson plan for the subject: "{subject}".

**Target Audience:** {difficulty_level} level students
**Overall Objective:** To provide a comprehensive and engaging learning experience that builds a strong foundation in {subject}, ensuring students grasp both the theoretical underpinnings and practical applications of each concept.


"""

    if parent_topics_content:
        prompt += "**Context from Parent Topics:**\n"
        for parent_id, parent_desc in parent_topics_content.items():
            prompt += f"  - **{parent_id}:** {parent_desc}\n"

    prompt += f"""
**Current Chunk:** {topic_details}

**Your Task:**
Generate detailed content for this specific chunk of the lesson plan. This is a part of a larger, cohesive plan, so maintain consistency in style, tone, and depth. Ensure that the content is engaging, informative, and suitable for in-depth learning.

"""

    # Depth-based instructions (refined)
    if depth == 1:
        prompt += "**Focus:** Provide a comprehensive overview, establishing the foundational concepts and clearly outlining the subtopics. Lay the groundwork for deeper exploration in subsequent chunks.\n"
    elif depth == 2:
        prompt += "**Focus:** Elaborate on the key concepts introduced earlier. Provide detailed explanations, incorporating examples and analogies to enhance understanding. Ensure a smooth transition from foundational concepts to more complex ideas.\n"
    elif depth >= 3:
        prompt += "**Focus:** Dive deep into the intricacies of each subtopic. Provide in-depth explanations, real-world applications, and challenging scenarios. Encourage critical thinking and problem-solving skills.\n"

    prompt += f"""
**Format and Content Requirements (Strictly Adhere to):**

{topic_details}


1. **Micro-Level Learning Objectives (3-5):** VERY IMPORTANT
    -   Define SMART (Specific, Measurable, Achievable, Relevant, Time-bound) objectives for this chunk.
    -   Begin each objective with an action verb (e.g., Define, Explain, Analyze, Design, Implement).
    -   Ensure alignment with the overall objective of the lesson plan.  
    -   Explain the "why" behind these concepts – their importance and relevance.
    -   Use analogies, metaphors, or real-world examples to enhance understanding are highly encouraged.
    -   **Crucially:** Address potential misconceptions proactively. Anticipate common misunderstandings and clarify them before they take root.
2.
    -   Identify potential challenges or misconceptions that students might encounter.
    -   **ELI5:** If a concept is particularly complex, suggest creating a simplified "Explain Like I'm 5" section in the lecture notes.

**Guiding Principles:**
-   **Clarity and Precision:** Use clear, concise language. Avoid jargon or overly complex sentences.
-   **Engagement:** Maintain an enthusiastic and encouraging tone.
-   **Continuity:** Ensure a smooth flow from previous chunks.
-   **No Repetition:** Refer back to concepts briefly if needed, but do not repeat detailed explanations.
-   **Markdown Formatting:** Use markdown for formatting (headings, lists, bold, italics). No unnecessary asterisks.

"""

    return prompt

def generate_lesson_plan_chunk(subject, difficulty_level, topic_data, parent_topics_content=None, depth=1, temperature=0.7):
    """
    Generates detailed content for a specific chunk of the lesson plan, using hierarchical context.
    """
    prompt = build_prompt_with_hierarchy(subject, difficulty_level, topic_data, parent_topics_content, depth)

    try:
        if prompt in generation_cache:
            st.success(f"Content for {topic_data['id']} found in cache!")
            return generation_cache[prompt]
        with st.spinner(f"Generating content for {topic_data['id']} (Depth: {depth})..."):
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=temperature,
                    max_output_tokens=500,
                )
            )
        # Convert to Markdown-like format
        markdown_content = response.text.strip()
        markdown_content = markdown_content.replace("*   ", "- ")  # Basic list conversion

        generation_cache[prompt] = markdown_content
        return markdown_content
    except Exception as e:
        st.error(f"Error generating lesson plan chunk: {e}")
        print(f"Error generating lesson plan chunk: {e}")  # Debugging
        return ""

def generate_lesson_plan_recursive(subject, roadmap, difficulty_level, temperature=0.7, parent_topics_content=None, depth=1):
    """
    Generates a comprehensive lesson plan recursively and returns a JSON structure, incorporating depth.
    """
    roadmap_dict = parse_roadmap(roadmap)
    lesson_plan_json = {
        "subject": subject,
        "difficulty": difficulty_level,
        "topics": []
    }

    for topic in roadmap_dict["topics"]:
        topic_json = generate_lesson_plan_chunk_json(
            subject, difficulty_level, topic, temperature, parent_topics_content, depth
        )
        lesson_plan_json["topics"].append(topic_json)

    return lesson_plan_json

def generate_lesson_plan_chunk_json(subject, difficulty_level, topic_data, temperature, parent_topics_content=None, depth=1):
    """
    Recursively generates content for a topic/subtopic and returns it as a JSON object, handling depth.
    Correctly handles arbitrary levels of nesting.
    """

    # Build parent_topics_content for subtopics
    current_level_context = {
        topic_data["id"]: topic_data["description"]
    }
    if parent_topics_content:
        current_level_context.update(parent_topics_content)

    content_string = generate_lesson_plan_chunk(
        subject, difficulty_level, topic_data, current_level_context, depth, temperature
    )

    topic_json = {
        "id": topic_data["id"],
        "title": topic_data["description"],
        "content": content_string,
    }

    # Check for and handle subtopics
    if "subtopics" in topic_data and topic_data["subtopics"]:
        topic_json["subtopics"] = []
        for subtopic in topic_data["subtopics"]:
            subtopic_json = generate_lesson_plan_chunk_json(
                subject, difficulty_level, subtopic, temperature, current_level_context, depth + 1
            )
            topic_json["subtopics"].append(subtopic_json)

    # Check for and handle subsubtopics
    if "subsubtopics" in topic_data and topic_data["subsubtopics"]:
        topic_json["subsubtopics"] = []
        for subsubtopic in topic_data["subsubtopics"]:
            subsubtopic_json = generate_lesson_plan_chunk_json(
                subject, difficulty_level, subsubtopic, temperature, current_level_context, depth + 2
            )
            topic_json["subsubtopics"].append(subsubtopic_json)

    # Check for and handle subsubsubtopics
    if "subsubsubtopics" in topic_data and topic_data["subsubsubtopics"]:
        topic_json["subsubsubtopics"] = []
        for subsubsubtopic in topic_data["subsubsubtopics"]:
            subsubsubtopic_json = generate_lesson_plan_chunk_json(
                subject, difficulty_level, subsubsubtopic, temperature, current_level_context, depth + 3
            )
            topic_json["subsubsubtopics"].append(subsubsubtopic_json)

    return topic_json

def save_lesson_plan_json(lesson_plan_json, filename="lesson_plan.json"):
    """Saves the lesson plan JSON to a file."""
    try:
        with open(filename, "w") as f:
            json.dump(lesson_plan_json, f, indent=4)
        st.success(f"Lesson plan saved as {filename}")
    except Exception as e:
        st.error(f"Error saving lesson plan: {e}")
        print(f"Error saving lesson plan: {e}")  # Debugging

def display_lesson_plan_for_editing(lesson_plan_json):
    """
    Displays the lesson plan from the JSON for editing in Streamlit.
    """
    for topic in lesson_plan_json["topics"]:
        display_topic(topic, level=1)

def display_topic(topic, level):
    """
    Displays a single topic or subtopic using markdown headings and text areas.
    """
    st.markdown(f"{'#' * level} {topic['id']}: {topic['title']}")

    content = st.text_area(
        "Content",
        value=topic["content"],
        height=300,
        key=f"{topic['id']}_content"
    )
    topic["content"] = content

    for subtopic in topic.get("subtopics", []):
        display_topic(subtopic, level + 1)

    for subsubtopic in topic.get("subsubtopics", []):
        display_topic(subsubtopic, level + 2)

    for subsubsubtopic in topic.get("subsubsubtopics", []):
        display_topic(subsubsubtopic, level + 3)



def create_docx_from_markdown(text, filename):
    """
    Creates a DOCX file from the given text, interpreting it as Markdown and applying appropriate formatting.
    """
    try:
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)

        # Enhanced Markdown parsing
        paragraphs = text.split('\n')
        in_list = False
        list_level = 0
        in_code_block = False

        for para in paragraphs:
            para = para.strip()

            # Code blocks
            if para.startswith("```"):
                in_code_block = not in_code_block
                if in_code_block:
                    # Add a paragraph for the code block
                    code_para = document.add_paragraph()
                    code_para.style = document.styles.add_style(f'CodeBlock{len(document.styles)}', 1)
                    code_para.style.font.name = 'Courier New'
                    code_para.style.font.size = Pt(10)
                continue

            if in_code_block:
                code_para.add_run(para + '\n')
                continue

            # Headings
            if para.startswith('#'):
                level = para.count('#')
                heading_text = para.lstrip('# ').strip()
                document.add_heading(heading_text, level=level)
                continue

            # Lists
            if para.startswith('- ') or para.startswith('* '):
                if not in_list:
                    in_list = True
                    list_level = 1
                else:
                    # Check for nested list
                    spaces = len(para) - len(para.lstrip())
                    new_level = spaces // 2 + 1
                    if new_level > list_level:
                        list_level = new_level
                    elif new_level < list_level:
                        list_level = new_level

                list_item = para.lstrip('-* ').strip()
                p = document.add_paragraph(list_item, style='List Bullet' if list_level == 1 else f'List Bullet {list_level}')
                if list_level > 1:
                    p.paragraph_format.left_indent = Inches(0.5 * list_level)
                continue
            elif in_list:
                in_list = False
                list_level = 0

            # Bold and italics
            while '**' in para:
                start = para.find('**')
                end = para.find('**', start + 2)
                if end == -1:
                    break
                bold_text = para[start+2:end]
                para = para[:start] + '<<BOLD>>' + bold_text + '<<BOLD>>' + para[end+2:]

            while '*' in para:
                start = para.find('*')
                end = para.find('*', start + 1)
                if end == -1:
                    break
                italic_text = para[start+1:end]
                para = para[:start] + '<<ITALIC>>' + italic_text + '<<ITALIC>>' + para[end+1:]

            if para:
                p = document.add_paragraph()
                segments = para.split('<<')
                for segment in segments:
                    if segment.startswith('BOLD>>'):
                        run = p.add_run(segment[6:-6])
                        run.bold = True
                    elif segment.startswith('ITALIC>>'):
                        run = p.add_run(segment[8:-8])
                        run.italic = True
                    else:
                        p.add_run(segment)

        document.save(filename)
        return filename
    except Exception as e:
        st.error(f"Error creating DOCX from Markdown: {e}")
        print(f"Error creating DOCX from Markdown: {e}")
        return None

def create_docx_from_lesson_plan(lesson_plan_json, filename):
    """Creates a DOCX file from the lesson plan JSON."""
    try:
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)

        def add_content(data, level):
            if isinstance(data, list):
                for item in data:
                    add_content(item, level)
            elif isinstance(data, dict):
                if "id" in data and "title" in data:
                    document.add_heading(f"{data['id']}: {data['title']}", level=level)
                if "content" in data:
                    #document.add_paragraph(data["content"])
                    paragraphs = data["content"].split('\n')
                    for para in paragraphs:
                        para = para.strip()
                        if para.startswith('# '):
                            document.add_heading(para[2:], level=1)
                        elif para.startswith('## '):
                            document.add_heading(para[3:], level=2)
                        elif para.startswith('### '):
                            document.add_heading(para[4:], level=3)
                        elif para.startswith('#### '):
                            document.add_heading(para[5:], level=4)
                        elif para.startswith('- '):
                            document.add_paragraph(para[2:], style='List Bullet')
                        elif para:
                            document.add_paragraph(para)
                add_content(data.get("subtopics", []), level + 1)
                add_content(data.get("subsubtopics", []), level + 2)
                add_content(data.get("subsubsubtopics", []), level + 3)

        add_content(lesson_plan_json["topics"], level=2)  # Start with level 2 headings

        document.save(filename)
        return filename
    except Exception as e:
        st.error(f"Error creating lesson plan DOCX: {e}")
        print(f"Error creating lesson plan DOCX: {e}")
        return None

def extract_lesson_plan_entry(lesson_plan_json, current_id):
    """
    Extracts the relevant section from the lesson plan JSON based on ID.
    """
    def find_entry_recursive(data, target_id):
        if isinstance(data, list):
            for item in data:
                result = find_entry_recursive(item, target_id)
                if result:
                    return result
        elif isinstance(data, dict):
            if "id" in data and data["id"] == target_id:
                return data["content"]
            for key, value in data.items():
                result = find_entry_recursive(value, target_id)
                if result:
                    return result
        return None

    return find_entry_recursive(lesson_plan_json, current_id)

def has_sub_chunks(lesson_plan_json, current_id):
    """
    Checks if a given ID in the lesson plan JSON has sub-chunks
    (subtopics, subsubtopics, etc.) at any level.
    """
    def find_entry_recursive(data, target_id):
        if isinstance(data, list):
            for item in data:
                result = find_entry_recursive(item, target_id)
                if result:
                    return result
        elif isinstance(data, dict):
            if "id" in data and data["id"] == target_id:
                return data
            for key, value in data.items():
                result = find_entry_recursive(value, target_id)
                if result:
                    return result
        return None

    def has_sub_keys_recursive(data):
        if isinstance(data, list):
            for item in data:
                if has_sub_keys_recursive(item):
                    return True
        elif isinstance(data, dict):
            if "subtopics" in data and data["subtopics"]:
                return True
            if "subsubtopics" in data and data["subsubtopics"]:
                return True
            if "subsubsubtopics" in data and data["subsubsubtopics"]:
                return True
            for key, value in data.items():
                if has_sub_keys_recursive(value):
                    return True
        return False

    entry = find_entry_recursive(lesson_plan_json["topics"], current_id)
    if entry:
        return has_sub_keys_recursive(entry)
    return False

def get_sub_chunks(lesson_plan_json, current_id):
    """
    Gets the sub-chunk IDs for a given ID.
    """
    def find_subtopic_ids_recursive(data, target_id, found=False):
        sub_chunks = []

        if isinstance(data, list):
            for item in data:
                sub_chunks.extend(find_subtopic_ids_recursive(item, target_id, found))
        elif isinstance(data, dict):
            if "id" in data and data["id"] == target_id:
                found = True

            if found:
                if "subtopics" in data:
                    for subtopic in data["subtopics"]:
                        sub_chunks.append(subtopic["id"])
                if "subsubtopics" in data:
                    for subsubtopic in data["subsubtopics"]:
                        sub_chunks.append(subsubtopic["id"])
                if "subsubsubtopics" in data:
                    for subsubsubtopic in data["subsubsubtopics"]:
                        sub_chunks.append(subsubsubtopic["id"])

            if not found or "subtopics" in data or "subsubtopics" in data or "subsubsubtopics" in data:
                for value in data.values():
                    sub_chunks.extend(find_subtopic_ids_recursive(value, target_id, found))

        return sub_chunks

    return find_subtopic_ids_recursive(lesson_plan_json["topics"], current_id)

def create_lecture_notes_prompt(lesson_plan_entry, current_id, difficulty_level, highlighted_topics, parent_topics_content=None):
    """
    Creates a highly optimized prompt for generating detailed lecture notes, incorporating parent topic context, specific instructions, advanced learning strategies, and addressing potential issues.
    """
    prompt = f"""
You are an expert educator creating comprehensive lecture notes for the following topic:

**Topic ID:** {current_id}
**Subject:** {st.session_state.subject}
**Target Audience:** {difficulty_level} level students
**Overall Objective:** To deliver a deep and engaging learning experience that equips students with a thorough understanding of {st.session_state.subject}, emphasizing both the theoretical foundations and practical applications of each concept.

**Contextual Reminders:**
- These lecture notes build upon the previously generated lesson plan.
- Ensure continuity and a logical flow, building upon knowledge established in previous sections.
- Avoid repetition. Briefly reference prior concepts if needed, but focus on new material.
- Maintain consistency in style and tone throughout the notes.

"""

    if parent_topics_content:
        prompt += "**Context from Parent Topics:**\n"
        for parent_id, parent_content in parent_topics_content.items():
            prompt += f"  - **{parent_id}:** {parent_content}\n"

    prompt += f"""
**Lesson Plan Context (Reference):**
{lesson_plan_entry}

**Highlighted Topics (for numericals/examples):**
{highlighted_topics}

**Your Task:**
Generate detailed lecture notes for this topic based on the provided lesson plan context. These notes should be comprehensive, engaging, and suitable for in-depth learning.

**Incorporate the following advanced learning strategies and content requirements:**

- Use clear, concise language appropriate for the target audience. Define all technical terms.
- Provide thorough explanations of each concept, principle, and process. 
-   **Illustrative Examples:** Incorporate numerous examples to illustrate concepts. Use a variety of examples(Conceptual, Real-World, Numerical examples for difficult topics)
-   **Analogies and Metaphors:** Use **analogies and metaphors** to explain complex concepts in simpler terms.
-   **Addressing Misconceptions:** Proactively address potential misconceptions identified in the lesson plan.
-   **Emphasis on Highlighted Topics:** Pay special attention to the topics highlighted for numericals. Provide extra examples and practice problems in these areas.
-   **ELI5 (Explain Like I'm 5):** For particularly complex concepts (as indicated in the lesson plan), include a simplified explanation using basic language and analogies, suitable for someone with no prior knowledge of the subject.
-   **Concept Maps:** Where appropriate, include or suggest the creation of **concept maps** that visually represent the relationships between key concepts.
-   **Markdown Formatting:** Use markdown for formatting (headings, subheadings, lists, bold, italics, code blocks). Ensure proper formatting and readability. No unnecessary use of asterisks.
-   **Interactive Elements:** Suggest interactive elements for the lecture at the last of each major section, such as:
    Think-Pair-Share, Quick Polls, Problem-Solving Breaks

**Guiding Principles:**
-   **Student-Centric:** Focus on student understanding and engagement.
-   **Comprehensive:** Cover all aspects of the topic in extreme detail.
-   **Engaging:** Use a conversational tone, ask questions, and encourage active learning.
-   **Continuity:** Maintain a strong connection with the lesson plan and previous topics.

    """

    return prompt

def generate_text_from_prompt(prompt, temperature=0.8):
    """Generates text from a prompt using the LLM."""
    try:
        if prompt in generation_cache:
            st.success("Text found in cache!")
            return generation_cache[prompt]
        with st.spinner("Generating text..."):
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=temperature,

                )
            )
        st.success("Text generated!")
        generation_cache[prompt] = response.text.strip()
        return response.text.strip()
    except Exception as e:
        st.error(f"Error generating text: {e}")
        print(f"Error generating text: {e}")
        return ""

def format_lecture_notes_content(content, current_id):
    """Formats the generated lecture notes content."""
    return content

def generate_lecture_notes_chunk(lesson_plan_json, current_id, difficulty_level, highlighted_topics, parent_topics_content=None, temperature=0.8):
    """Generates lecture notes for a specific chunk, without recursively processing sub-chunks."""
    lesson_plan_entry = extract_lesson_plan_entry(lesson_plan_json, current_id)

    # Build parent_topics_context for subtopics
    current_level_context = {}
    if parent_topics_content:
        current_level_context.update(parent_topics_content)

    # Generate content for the current topic only
    prompt = create_lecture_notes_prompt(lesson_plan_entry, current_id, difficulty_level, highlighted_topics, current_level_context)
    generated_content = generate_text_from_prompt(prompt, temperature)
    formatted_content = format_lecture_notes_content(generated_content, current_id)

    return formatted_content  # Return only the content for the current topic

def create_detailed_notes_recursive(lesson_plan_json, difficulty_level, highlighted_topics, temperature=0.8):
    """
    Generates detailed lecture notes recursively, ensuring all topics and subtopics are covered.
    """
    document_text = ""
    processed_ids = set()  # Keep track of processed IDs to prevent duplication

    def count_items_recursive(data):
        count = 0
        if isinstance(data, list):
            for item in data:
                count += count_items_recursive(item)
        elif isinstance(data, dict):
            if "id" in data:
                count += 1
            for key, value in data.items():
                if key in ["subtopics", "subsubtopics", "subsubsubtopics"]:
                    count += count_items_recursive(value)
        return count

    total_items = count_items_recursive(lesson_plan_json["topics"])
    item_count = 0
    progress_bar = st.progress(0)

    def generate_notes_recursive(data, parent_topics_context=None):
        nonlocal document_text, item_count

        if isinstance(data, list):
            for item in data:
                generate_notes_recursive(item, parent_topics_context)
        elif isinstance(data, dict):
            if "id" in data:
                topic_id = data["id"]

                if topic_id in processed_ids:
                    return  # Skip if already processed

                processed_ids.add(topic_id)  # Mark this topic as processed

                # Build parent_topics_context for subtopics
                current_level_context = {}
                if parent_topics_context:
                    current_level_context.update(parent_topics_context)
                current_level_context[topic_id] = data["title"]

                # Always call generate_lecture_notes_chunk for each topic
                topic_content = generate_lecture_notes_chunk(
                    lesson_plan_json, topic_id, difficulty_level, highlighted_topics, current_level_context, temperature
                )
                document_text += topic_content + "\n\n"

                item_count += 1
                progress_bar.progress(item_count / total_items if total_items > 0 else 1.0)

                # Process subtopics immediately after each topic
                for key in ["subtopics", "subsubtopics", "subsubsubtopics"]:
                    if key in data:
                        generate_notes_recursive(data[key], current_level_context)

    generate_notes_recursive(lesson_plan_json["topics"], None)
    filename = "detailed_notes.docx"
    if create_docx_from_markdown(document_text, filename):
        st.success(f"Detailed notes saved as {filename}")
        return filename
    else:
        return None

# --- Streamlit UI ---
st.markdown("""
<style>
.stProgress > div > div > div > div {
    background-color: #4CAF50;
    height: 10px;
}
</style>""", unsafe_allow_html=True)

if 'name' not in st.session_state:
    st.session_state.name = "User"
st.sidebar.write(f'Welcome *{st.session_state.name}*')
st.title("Syllabus to Detailed Notes Generator")

st.header("Step 1: Syllabus and Difficulty")
uploaded_syllabus = st.file_uploader("Upload Syllabus Text File (.txt)", type=["txt"])
difficulty_level = st.selectbox("Select Difficulty Level", ["Btech", "Mtech", "PHD"])
subject = st.text_input("Enter the subject name:")
st.session_state.subject = subject

st.header("Step 2: Generate and Edit Roadmap")
if uploaded_syllabus:
    syllabus_text = uploaded_syllabus.read().decode("utf-8")
    if st.button("Generate Roadmap"):
        roadmap = generate_roadmap(subject, syllabus_text, difficulty_level)
        st.session_state.roadmap = roadmap

if "roadmap" in st.session_state:
    edited_roadmap_text = st.text_area("Edit Roadmap", value=st.session_state.roadmap, height=300)
    if st.button("Save Edited Roadmap"):
        st.session_state.roadmap = edited_roadmap_text
        st.success("Roadmap updated!")

st.header("Step 3: Generate and Edit Lesson Plan")
llm_temperature_plan = st.slider("LLM Temperature (Plan)", 0.0, 1.0, 0.7, step=0.1,
                                    help="Controls randomness of lesson plan generation.")
depth = st.select_slider(
    "Select Depth of Lesson Plan",
    options=[1, 2, 3],
    format_func=lambda x: {1: "Low", 2: "Medium", 3: "High"}[x],
    help="Controls the level of detail in the generated lesson plan."
)
if "roadmap" in st.session_state:
    if st.button("Generate Lesson Plan"):
        with st.spinner("Generating lesson plan..."):
            lesson_plan_json = generate_lesson_plan_recursive(
                subject, st.session_state.roadmap, difficulty_level, llm_temperature_plan, None, depth
            )
            st.session_state.lesson_plan = lesson_plan_json
            save_lesson_plan_json(lesson_plan_json)
        st.success("Lesson plan generated and saved as JSON!")

if "lesson_plan" in st.session_state:
    st.header("Edit Lesson Plan")
    display_lesson_plan_for_editing(st.session_state.lesson_plan)

    if st.button("Save Edited Lesson Plan"):
        save_lesson_plan_json(st.session_state.lesson_plan)
        st.success("Edited lesson plan saved!")

    if st.button("Download Edited Lesson Plan as docx"):
        create_docx_from_lesson_plan(st.session_state.lesson_plan, "edited_lesson_plan.docx")

st.header("Step 4: Generate Detailed Notes")
highlighted_topics_input = st.text_area("Highlight topics important for numericals (comma-separated):", "")
llm_temperature_notes = st.slider("LLM Temperature (Notes)", 0.0, 1.0, 0.8, step=0.1,
                                   help="Controls randomness of notes generation.")

if "lesson_plan" in st.session_state:
    if st.button("Generate Detailed Notes"):
        highlighted_topics = [t.strip() for t in highlighted_topics_input.split(",") if t.strip()]
        with st.spinner("Generating detailed notes..."):
            notes_filename = create_detailed_notes_recursive(
                st.session_state.lesson_plan, difficulty_level, highlighted_topics, llm_temperature_notes
            )
            if notes_filename:  # Check if notes
                st.session_state.notes_filename = notes_filename

    # Download button only if notes have been generated
    if "notes_filename" in st.session_state:
        with open(st.session_state.notes_filename, "rb") as f:
            st.download_button(
                label="Download Detailed Notes (DOCX)",
                data=f,
                file_name=st.session_state.notes_filename
            )
        st.success("Detailed notes generated and ready for download!")
    elif 'lesson_plan' in st.session_state:
        st.warning("Please generate detailed notes first.")
else:
    st.warning("Please generate and save the lesson plan first.")
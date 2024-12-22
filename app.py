import streamlit as st
from docx import Document
from docx.shared import Pt
import os
from openai import OpenAI
from embedchain import App as EmbedchainApp
import io
from pypdf import PdfReader
import re
from streamlit_authenticator import Authenticate
import yaml
from yaml.loader import SafeLoader
import time
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# --- Setup ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    st.error("Please set your OpenAI API key in the .env file.")
    st.stop()

client = OpenAI(api_key=OPENAI_API_KEY)

# --- Authentication ---
try:
    with open('./config.yaml') as file:
        config = yaml.load(file, Loader=SafeLoader)
except FileNotFoundError:
    st.error("Error: config.yaml not found. Please create one based on the example.")
    st.stop()
except yaml.YAMLError as e:
    st.error(f"Error parsing config.yaml: {e}")
    st.stop()

authenticator = Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days'],
    config['preauthorized']
)

# Check for preauthorized users
if 'name' not in st.session_state:
    st.session_state.name = None
if 'authentication_status' not in st.session_state:
    st.session_state.authentication_status = None
if 'username' not in st.session_state:
    st.session_state.username = None

if st.session_state.authentication_status is None or st.session_state.authentication_status is False:
    if st.session_state.username and st.session_state.username in config['preauthorized']['emails']:
        # Preauthorized user logic
        st.session_state.authentication_status = True
        st.session_state.name = config['credentials']['usernames'][st.session_state.username]['name']
    else:
        # Prompt for login if not preauthorized
        name, authentication_status, username = authenticator.login('Login', 'sidebar')
        st.session_state.name = name
        st.session_state.authentication_status = authentication_status
        st.session_state.username = username

        if st.session_state.authentication_status is False:
            st.error('Username/password is incorrect')
        elif st.session_state.authentication_status is None:
            st.warning('Please enter your username and password')

# Main app logic only runs if authenticated or preauthorized
if st.session_state.authentication_status:
    if st.session_state.username not in config['preauthorized']['emails']:
        # Regular logout for users who authenticated via login form
        authenticator.logout('Logout', 'sidebar')
    st.sidebar.write(f'Welcome *{st.session_state.name}*')

    # --- Helper Functions ---
    # (These are the same as before - no changes needed here)
    def generate_lesson_plan(syllabus_text, difficulty_level, temperature=0.7, max_tokens=1000):
        prompt = f"""Generate a detailed lesson plan based on the following syllabus and difficulty level:
            Syllabus: {syllabus_text}
            Difficulty Level: {difficulty_level}

            The lesson plan should be structured in a hierarchical format, using "->" to denote different levels. For example:
            Topic 1
            Topic 1 -> Subtopic 1
            Topic 1 -> Subtopic 1 -> Sub-subtopic 1
            Topic 2
            Topic 2 -> Subtopic 1

            Ensure the plan is comprehensive and covers all aspects of the syllabus at the specified difficulty level."""
        try:
            response = client.completions.create(model="gpt-3.5-turbo-instruct", prompt=prompt, max_tokens=max_tokens,
                                                temperature=temperature)
            return response.choices[0].text.strip()
        except Exception as e:
            st.error(f"Error generating lesson plan: {e}")
            return None

    def create_docx_from_text(text, filename="lesson_plan.docx"):
        document = Document()
        document.add_paragraph(text)
        document.save(filename)
        return filename

    def generate_slide_content(topic, subtopic, sub_subtopic, difficulty_level, pdf_content=None, temperature=0.8,
                               max_tokens=1500):
        prompt = f"""Generate detailed content for a slide on the topic: '{topic}', subtopic: '{subtopic}', sub-subtopic: '{sub_subtopic}' for a {difficulty_level} level course. 
            Include relevant explanations, examples, and details suitable for comprehensive understanding."""
        context_str = f"\nConsider the following information:\n{pdf_content}" if pdf_content else ""
        try:
            response = client.completions.create(model="gpt-3.5-turbo-instruct", prompt=prompt + context_str,
                                                max_tokens=max_tokens, temperature=temperature)
            return response.choices[0].text.strip()
        except Exception as e:
            st.error(f"Error generating slide content: {e}")
            return None

    def generate_image_prompt(topic, subtopic, sub_subtopic, temperature=0.7, max_tokens=100):
        prompt = f"Generate a descriptive prompt for an image representing the topic: '{topic}', subtopic: '{subtopic}', sub-subtopic: '{sub_subtopic}'."
        try:
            response = client.completions.create(model="gpt-3.5-turbo-instruct", prompt=prompt, max_tokens=max_tokens,
                                                temperature=temperature)
            return response.choices[0].text.strip()
        except Exception as e:
            st.warning(f"Could not generate image prompt: {e}")
            return None

    def create_detailed_notes_docx(lesson_plan_structure, difficulty_level, pdf_content=None):
        document = Document()
        document.add_heading("Detailed Course Notes", level=1)
        total_slides = sum(
            len(subtopics.get(subtopic, [])) if isinstance(subtopics, dict) else 1 for topic, subtopics in
            lesson_plan_structure.items())
        progress_bar = st.progress(0)
        slide_count = 0
        for topic, subtopics in lesson_plan_structure.items():
            with st.expander(f"Topic: {topic}", expanded=True):
                st.info(f"Generating content for: {topic}")
                document.add_heading(topic, level=2)
                if isinstance(subtopics, dict):
                    for subtopic, sub_subtopics in subtopics.items():
                        document.add_heading(subtopic, level=3)
                        for sub_subtopic in sub_subtopics:
                            slide_content = generate_slide_content(topic, subtopic, sub_subtopic, difficulty_level,
                                                                   pdf_content)
                            if slide_content:
                                document.add_heading(sub_subtopic, level=4)
                                paragraph = document.add_paragraph(slide_content)
                                paragraph_format = paragraph.paragraph_format
                                paragraph_format.space_before = Pt(6)
                                paragraph_format.space_after = Pt(6)
                                image_prompt = generate_image_prompt(topic, subtopic, sub_subtopic)
                                if image_prompt:
                                    document.add_paragraph(f"**Image Prompt:** {image_prompt}", style='Intense Quote')
                            slide_count += 1
                            progress_bar.progress(slide_count / total_slides if total_slides > 0 else 1.0)
                else:
                    slide_content = generate_slide_content(topic, "", "", difficulty_level, pdf_content)
                    if slide_content:
                        document.add_paragraph(slide_content)
                    slide_count += 1
                    progress_bar.progress(slide_count / total_slides if total_slides > 0 else 1.0)
        filename = "detailed_notes.docx"
        document.save(filename)
        return filename

    def extract_text_from_pdf(pdf_file):
        text = ""
        try:
            pdf_reader = PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text += page.extract_text()
        except Exception as e:
            st.error(f"Error extracting text from PDF: {e}")
        return text

    def parse_lesson_plan(lesson_plan_text):
        lesson_plan_structure = {}
        current_topic = None
        current_subtopic = None
        for line in lesson_plan_text.strip().split('\n'):
            line = line.strip()
            if not line: continue
            if "->" not in line:
                current_topic = line
                lesson_plan_structure[current_topic] = {}
                current_subtopic = None
            elif line.count("->") == 1:
                topic, subtopic = map(str.strip, line.split("->"))
                if topic == current_topic:
                    lesson_plan_structure[current_topic][subtopic] = []
                    current_subtopic = subtopic
            elif line.count("->") == 2:
                topic, subtopic, sub_subtopic = map(str.strip, line.split("->"))
                if topic == current_topic and subtopic == current_subtopic:
                    lesson_plan_structure[current_topic][subtopic].append(sub_subtopic)
        return lesson_plan_structure
    # --- Streamlit UI ---
    # App logic moved inside the authentication check
    st.title("Syllabus to Detailed Notes Generator")

    # --- Step 1: Upload Syllabus and Select Difficulty ---
    st.header("Step 1: Syllabus and Difficulty")
    uploaded_syllabus = st.file_uploader("Upload Syllabus Text File (.txt)", type=["txt"])
    difficulty_level = st.selectbox("Select Difficulty Level", ["Btech", "Mtech", "PHD"])

    # --- Step 2: Generate and Edit Lesson Plan ---
    st.header("Step 2: Generate and Edit Lesson Plan")
    llm_temperature_plan = st.slider("LLM Temperature (Plan)", 0.0, 1.0, 0.7, step=0.1,
                                      help="Controls randomness of lesson plan generation.")
    llm_max_tokens_plan = st.slider("LLM Max Tokens (Plan)", 500, 2000, 1000, step=100,
                                    help="Maximum length of the generated lesson plan.")
    if uploaded_syllabus:
        syllabus_text = uploaded_syllabus.read().decode("utf-8")
        if st.button("Generate Initial Lesson Plan"):
            with st.spinner("Generating initial lesson plan..."):
                initial_lesson_plan = generate_lesson_plan(syllabus_text, difficulty_level, llm_temperature_plan,
                                                            llm_max_tokens_plan)
                st.session_state.lesson_plan_text = initial_lesson_plan if initial_lesson_plan else ""
    else:
        st.warning("Please upload a syllabus file to generate a lesson plan.")

    edited_lesson_plan_text = st.text_area("Edit Lesson Plan",
                                            value=st.session_state.get("lesson_plan_text", ""), height=300)

    if st.button("Save Edited Lesson Plan"):
        try:
            st.session_state.lesson_plan_structure = parse_lesson_plan(edited_lesson_plan_text)
            lesson_plan_filename = create_docx_from_text(edited_lesson_plan_text, "final_lesson_plan.docx")
            with open(lesson_plan_filename, "rb") as f:
                st.download_button(label="Download Lesson Plan (DOCX)", data=f, file_name=lesson_plan_filename)
            st.success("Lesson plan saved and parsed!")
        except Exception as e:
            st.error(f"Error parsing or saving lesson plan: {e}")

    # --- Step 3: Upload Optional Textbook and Generate Notes ---
    st.header("Step 3: Generate Detailed Notes")
    uploaded_book = st.file_uploader("Optional: Upload Textbook PDF", type=["pdf"])
    llm_temperature_notes = st.slider("LLM Temperature (Notes)", 0.0, 1.0, 0.8, step=0.1,
                                       help="Controls randomness of notes generation.")
    llm_max_tokens_notes = st.slider("LLM Max Tokens (Notes)", 500, 2500, 1500, step=100,
                                     help="Maximum length of the generated notes for each section.")

    if st.session_state.get("lesson_plan_structure"):
        if st.button("Generate Detailed Notes"):
            pdf_content = None
            if uploaded_book:
                with st.spinner("Extracting text from PDF..."):
                    pdf_content = extract_text_from_pdf(uploaded_book)
            with st.spinner("Generating detailed notes..."):
                notes_filename = create_detailed_notes_docx(st.session_state.lesson_plan_structure,
                                                            difficulty_level, pdf_content)
                st.session_state.notes_filename = notes_filename
                with open(notes_filename, "rb") as f:
                    st.download_button(label="Download Detailed Notes (DOCX)", data=f, file_name=notes_filename)
                st.success("Detailed notes generated and ready for download!")
    else:
        st.warning("Please generate and save the lesson plan first.")

    # --- Step 4: Ask Questions from Notes ---
    st.header("Step 4: Ask Questions from Notes")
    if st.session_state.get("notes_filename"):
        try:
            with open(st.session_state.notes_filename, "r", encoding="utf-8", errors='ignore') as f:
                notes_content = f.read()
                st.session_state.notes_content = notes_content  # Store for highlighting
        except Exception as e:
            notes_content = "Could not load notes for question answering."
            st.error(f"Error loading notes: {e}")

        notes_content_display = st.text_area("Review Notes",
                                              value=st.session_state.notes_content if 'notes_content' in st.session_state else "Generate notes to view.",
                                              height=200)
        question = st.text_input("Ask a question about the notes:")
        llm_temperature_qa = st.slider("LLM Temperature (QA)", 0.0, 1.0, 0.7, step=0.1,
                                       help="Controls randomness of question answering.")
        llm_max_tokens_qa = st.slider("LLM Max Tokens (QA)", 100, 1000, 500, step=50,
                                      help="Maximum length of the answer.")

        if st.button("Get Answer"):
            if question and 'notes_content' in st.session_state:
                with st.spinner("Getting answer..."):
                    prompt = f"Answer the following question based on the notes:\n\nNotes:\n{st.session_state.notes_content}\n\nQuestion: {question}"
                    try:
                        response = client.completions.create(model="gpt-3.5-turbo-instruct", prompt=prompt,
                                                            max_tokens=llm_max_tokens_qa,
                                                            temperature=llm_temperature_qa)
                        st.write("Answer:")
                        st.write(response.choices[0].text.strip())
                    except Exception as e:
                        st.error(f"Error answering question: {e}")
            elif not 'notes_content' in st.session_state:
                st.warning("Please generate detailed notes first.")
            else:
                st.warning("Please enter your question.")
    else:
        st.warning("Please generate detailed notes first to ask questions.")

else:
    st.write("Please log in to access the app.")
